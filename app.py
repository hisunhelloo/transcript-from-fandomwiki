"""
app.py  –  Flask web server for Fandom Transcript Crawler
"""

import io
import re
import urllib.request
import zipfile
from pathlib import Path
from urllib.parse import urlparse, urlencode, unquote

import requests
from bs4 import BeautifulSoup, NavigableString, Tag
from flask import Flask, jsonify, render_template, request, send_file

app = Flask(__name__)

# ---------------------------------------------------------------------------
# Transcript core logic  (same as crawl_transcript.py)
# ---------------------------------------------------------------------------

HEADERS = {
    "User-Agent": (
        "TranscriptCrawler/1.0 (educational; "
        "+https://github.com/user/transcript-from-fandom)"
    ),
    "Accept": "application/json",
}


class Line:
    DIALOGUE  = "dialogue"
    DIRECTION = "direction"
    SONG      = "song"

    def __init__(self, kind, character, text):
        self.kind      = kind
        self.character = character
        self.text      = text.strip()


def _wiki_api_url(fandom_url):
    parsed   = urlparse(fandom_url)
    api_base = f"{parsed.scheme}://{parsed.netloc}/api.php"
    # unquote() decodes percent-encoded chars: %27 → ', %28 → (, etc.
    title    = unquote(re.sub(r"^/wiki/", "", parsed.path))
    return api_base, title


def fetch_soup(url):
    api_url, page_title = _wiki_api_url(url)
    params = {
        "action": "parse",
        "page":   page_title,
        "prop":   "text",
        "format": "json",
        "disablelimitreport": 1,
    }
    full_url = f"{api_url}?{urlencode(params)}"
    resp = requests.get(full_url, headers=HEADERS, timeout=30)
    resp.raise_for_status()
    data = resp.json()
    if "error" in data:
        raise RuntimeError(f"MediaWiki API error: {data['error']}")
    return BeautifulSoup(data["parse"]["text"]["*"], "lxml")


def get_title(url):
    _, title = _wiki_api_url(url)
    return title.replace("_", " ").replace("/", " - ")


def _split_p_by_br(p_tag):
    rows, current = [], []
    for child in p_tag.children:
        if isinstance(child, Tag) and child.name == "br":
            rows.append(current)
            current = []
        else:
            current.append(child)
    if current:
        rows.append(current)
    return rows


def _nodes_to_text(nodes):
    parts = []
    for n in nodes:
        if isinstance(n, NavigableString):
            parts.append(str(n))
        elif isinstance(n, Tag):
            parts.append(n.get_text())
    return re.sub(r"\s+", " ", "".join(parts)).strip()


def _classify_row(nodes):
    if not nodes:
        return None
    raw = _nodes_to_text(nodes)
    if not raw or len(raw) < 2:
        return None

    first_tag = next(
        (n for n in nodes if isinstance(n, Tag) and n.name not in ("br",)), None
    )

    if first_tag and first_tag.name == "b":
        char = first_tag.get_text(strip=True).rstrip(":")
        after = []
        past  = False
        for n in nodes:
            if n is first_tag:
                past = True
                continue
            if past:
                after.append(n)
        dialogue = _nodes_to_text(after).lstrip(": ").strip()
        return Line(Line.DIALOGUE, char, dialogue)

    if all(
        isinstance(n, NavigableString) or (isinstance(n, Tag) and n.name == "i")
        for n in nodes
        if not (isinstance(n, NavigableString) and n.strip() == "")
    ):
        if raw.startswith("(") or (first_tag and first_tag.name == "i"):
            return Line(Line.DIRECTION, None, raw)

    if raw.startswith(("\u266a", "\u266b", "Song:", "\u2665")):
        return Line(Line.SONG, None, raw)

    if first_tag and first_tag.name == "i":
        return Line(Line.DIRECTION, None, raw)

    return Line(Line.DIRECTION, None, raw)


def parse_transcript(soup):
    content = soup.find("div", class_="mw-parser-output") or soup
    lines   = []
    for p in content.find_all("p", recursive=False):
        for row in _split_p_by_br(p):
            line = _classify_row(row)
            if line and len(line.text) > 1:
                lines.append(line)
    return lines


# ---------------------------------------------------------------------------
# Exporters
# ---------------------------------------------------------------------------

def to_markdown(title, url, lines):
    parts = [f"# {title}", "", f"> Source: <{url}>", "", "---", ""]
    for line in lines:
        if line.kind == Line.DIALOGUE:
            parts.append(f"**{line.character}:** {line.text}")
        elif line.kind == Line.SONG:
            parts.append(f"> *{line.text}*")
        else:
            parts.append(f"*{line.text}*")
        parts.append("")
    return "\n".join(parts)


# --- Font helpers for PDF ---
FONT_DIR     = Path(__file__).parent / "_fonts"
FONT_REGULAR = FONT_DIR / "DejaVuSans.ttf"
FONT_BOLD    = FONT_DIR / "DejaVuSans-Bold.ttf"
FONT_ITALIC  = FONT_DIR / "DejaVuSans-Oblique.ttf"
FONT_URLS    = {
    FONT_REGULAR: "https://cdn.jsdelivr.net/npm/dejavu-fonts-ttf@2.37.3/ttf/DejaVuSans.ttf",
    FONT_BOLD:    "https://cdn.jsdelivr.net/npm/dejavu-fonts-ttf@2.37.3/ttf/DejaVuSans-Bold.ttf",
    FONT_ITALIC:  "https://cdn.jsdelivr.net/npm/dejavu-fonts-ttf@2.37.3/ttf/DejaVuSans-Oblique.ttf",
}


def ensure_fonts():
    FONT_DIR.mkdir(parents=True, exist_ok=True)
    for dest, src_url in FONT_URLS.items():
        if not dest.exists():
            urllib.request.urlretrieve(src_url, dest)


def to_pdf_bytes(title, url, lines):
    ensure_fonts()
    from fpdf import FPDF

    class PDF(FPDF):
        def header(self):
            self.set_font("DejaVu", "B", 14)
            self.set_text_color(30, 30, 30)
            self.cell(0, 10, title, align="C", new_x="LMARGIN", new_y="NEXT")
            self.set_font("DejaVu", "I", 8)
            self.set_text_color(130, 130, 130)
            short = url if len(url) <= 90 else url[:87] + "..."
            self.cell(0, 5, f"Source: {short}", align="C", new_x="LMARGIN", new_y="NEXT")
            self.set_draw_color(200, 200, 200)
            self.set_line_width(0.3)
            self.ln(2)
            self.line(self.l_margin, self.get_y(), self.w - self.r_margin, self.get_y())
            self.ln(3)
            self.set_text_color(0, 0, 0)

        def footer(self):
            self.set_y(-15)
            self.set_font("DejaVu", "I", 8)
            self.set_text_color(160, 160, 160)
            self.cell(0, 10, f"Page {self.page_no()}", align="C")

    pdf = PDF(orientation="P", unit="mm", format="A4")
    pdf.add_font("DejaVu", "",  str(FONT_REGULAR))
    pdf.add_font("DejaVu", "B", str(FONT_BOLD))
    pdf.add_font("DejaVu", "I", str(FONT_ITALIC))
    pdf.set_auto_page_break(auto=True, margin=20)
    pdf.add_page()
    W = pdf.w - pdf.l_margin - pdf.r_margin

    for line in lines:
        if line.kind == Line.DIALOGUE:
            pdf.set_font("DejaVu", "B", 10)
            pdf.set_text_color(25, 75, 160)
            label = f"{line.character}: "
            lw    = min(pdf.get_string_width(label) + 1, W * 0.35)
            pdf.set_x(pdf.l_margin)
            pdf.cell(lw, 6, label, new_x="RIGHT", new_y="LAST")
            pdf.set_font("DejaVu", "", 10)
            pdf.set_text_color(20, 20, 20)
            pdf.multi_cell(W - lw, 6, line.text)
        elif line.kind == Line.SONG:
            pdf.set_font("DejaVu", "I", 10)
            pdf.set_text_color(100, 60, 160)
            pdf.set_x(pdf.l_margin + 5)
            pdf.multi_cell(W - 5, 5, line.text)
        else:
            pdf.set_font("DejaVu", "I", 9)
            pdf.set_text_color(90, 90, 90)
            pdf.multi_cell(W, 5, line.text)
        pdf.ln(1)

    buf = io.BytesIO()
    pdf.output(buf)
    return buf.getvalue()


def to_docx_bytes(title, url, lines):
    from docx import Document
    from docx.shared import Pt, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH

    doc = Document()

    # Title
    t = doc.add_heading(title, level=1)
    t.alignment = WD_ALIGN_PARAGRAPH.CENTER

    # Source
    src = doc.add_paragraph()
    src.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = src.add_run(f"Source: {url}")
    run.font.size   = Pt(8)
    run.font.italic = True
    run.font.color.rgb = RGBColor(0x88, 0x88, 0x88)

    doc.add_paragraph()  # spacer

    for line in lines:
        p = doc.add_paragraph()
        if line.kind == Line.DIALOGUE:
            r1 = p.add_run(f"{line.character}: ")
            r1.bold            = True
            r1.font.color.rgb  = RGBColor(0x19, 0x4B, 0xA0)
            r2 = p.add_run(line.text)
            r2.font.size = Pt(10)
        elif line.kind == Line.SONG:
            r = p.add_run(line.text)
            r.italic           = True
            r.font.color.rgb   = RGBColor(0x64, 0x3C, 0xA0)
            r.font.size        = Pt(10)
        else:
            r = p.add_run(line.text)
            r.italic           = True
            r.font.color.rgb   = RGBColor(0x55, 0x55, 0x55)
            r.font.size        = Pt(9)

    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()


# ---------------------------------------------------------------------------
# Routes
# ---------------------------------------------------------------------------

@app.route("/")
def index():
    return render_template("index.html")


@app.route("/api/generate", methods=["POST"])
def generate():
    data    = request.get_json()
    url     = (data.get("url") or "").strip()
    formats = data.get("formats", [])   # list of "md", "pdf", "docs"

    if not url:
        return jsonify({"error": "URL is required."}), 400
    if not formats:
        return jsonify({"error": "Select at least one format."}), 400

    # Validate: must look like a Fandom /Transcript URL
    if "fandom.com" not in url:
        return jsonify({"error": "Please enter a valid Fandom wiki URL."}), 400

    try:
        soup  = fetch_soup(url)
        title = get_title(url)
        lines = parse_transcript(soup)

        if not lines:
            return jsonify({"error": "No transcript content found on this page."}), 422

        # Build slug for filenames
        path_part = urlparse(url).path.strip("/").replace("/", "_")
        slug = re.sub(r"^wiki_", "", path_part)
        slug = re.sub(r"[^\w]", "_", slug)

        # Pack into a zip
        zip_buf = io.BytesIO()
        with zipfile.ZipFile(zip_buf, "w", zipfile.ZIP_DEFLATED) as zf:
            if "md" in formats:
                content = to_markdown(title, url, lines)
                zf.writestr(f"{slug}.md", content.encode("utf-8"))

            if "pdf" in formats:
                zf.writestr(f"{slug}.pdf", to_pdf_bytes(title, url, lines))

            if "docs" in formats:
                zf.writestr(f"{slug}.docx", to_docx_bytes(title, url, lines))

        zip_buf.seek(0)
        return send_file(
            zip_buf,
            mimetype="application/zip",
            as_attachment=True,
            download_name=f"{slug}.zip",
        )

    except requests.HTTPError as e:
        return jsonify({"error": f"Failed to fetch page: {e}"}), 502
    except Exception as e:
        return jsonify({"error": str(e)}), 500


if __name__ == "__main__":
    import os
    port  = int(os.environ.get("PORT", 5000))
    debug = os.environ.get("FLASK_ENV") == "development"
    app.run(debug=debug, host="0.0.0.0", port=port)
